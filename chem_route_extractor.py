#!/usr/bin/env python
"""
ChemRouteExtractor - 从化学文献PDF中自动提取反应合成路线的工具

基于OpenChemIE，能够从化学文献图表中提取化学反应，并生成包含反应图像和实验步骤的Word文档。

用法示例：
    python chem_route_extractor.py --input literature.pdf --output ./results
    python chem_route_extractor.py --input ./pdfs --output ./results --max-pages 10
    python chem_route_extractor.py --input literature.pdf --no-images --language en
"""

import argparse
import sys
import os
from pathlib import Path
import logging
import json
from datetime import datetime
from typing import List, Dict, Any, Optional

# 尝试导入pandas（用于更好的时间戳格式）
try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def setup_argparse() -> argparse.ArgumentParser:
    """设置命令行参数解析"""
    parser = argparse.ArgumentParser(
        description='从化学文献PDF中自动提取反应合成路线的工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s --input literature.pdf --output ./results
  %(prog)s --input ./pdfs --output ./results --max-pages 10
  %(prog)s --input literature.pdf --no-images --language en
        """
    )
    
    parser.add_argument(
        '--input', '-i',
        required=True,
        help='输入PDF文件或包含PDF文件的目录'
    )
    
    parser.add_argument(
        '--output', '-o',
        default='./results',
        help='输出目录（默认: ./results）'
    )
    
    parser.add_argument(
        '--max-pages',
        type=int,
        default=5,
        help='每个PDF处理的最大页数（默认: 5）'
    )
    
    parser.add_argument(
        '--no-images',
        action='store_true',
        help='不生成反应图像'
    )
    
    parser.add_argument(
        '--include-si',
        action='store_true',
        help='同时处理SI（补充信息）文件'
    )
    
    parser.add_argument(
        '--language',
        choices=['zh', 'en', 'both'],
        default='both',
        help='输出语言：zh=中文，en=英文，both=中英双语（默认）'
    )
    
    parser.add_argument(
        '--template',
        help='自定义Word模板文档路径（可选）'
    )
    
    parser.add_argument(
        '--config',
        help='JSON配置文件路径（可选）'
    )
    
    parser.add_argument(
        '--debug',
        action='store_true',
        help='启用调试模式'
    )
    
    return parser

def load_config(config_path: str) -> Dict[str, Any]:
    """加载JSON配置文件"""
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        logger.error(f"加载配置文件失败: {e}")
        return {}

def find_pdf_files(input_path: str, include_si: bool = False) -> List[Path]:
    """查找PDF文件"""
    input_path = Path(input_path)
    pdf_files = []
    
    if input_path.is_file():
        if input_path.suffix.lower() == '.pdf':
            pdf_files.append(input_path)
        else:
            logger.error(f"输入文件不是PDF: {input_path}")
    elif input_path.is_dir():
        # 查找目录中的所有PDF文件
        for pdf_file in input_path.glob("**/*.pdf"):
            if not include_si and 'SI' in pdf_file.name:
                continue
            pdf_files.append(pdf_file)
    else:
        logger.error(f"输入路径不存在: {input_path}")
    
    return pdf_files

def extract_text_info(pdf_path: Path, max_pages: int = 5) -> Dict[str, Any]:
    """从PDF提取文本信息（实验部分、化合物等）"""
    import pdfplumber
    import re
    
    logger.info(f"提取文本信息: {pdf_path.name}")
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            pages_to_process = min(max_pages, total_pages)
            
            # 提取所有文本
            all_text = ""
            for i in range(pages_to_process):
                page = pdf.pages[i]
                text = page.extract_text()
                if text:
                    all_text += text + "\n\n"
            
            # 查找实验部分
            experimental_text = ""
            exp_patterns = [
                r'(?:Experimental|EXPERIMENTAL|Materials and Methods|METHODS)(.*?)(?=\n\n[A-Z]|$)',
                r'(?:Synthesis|SYNTHESIS|Preparation|PREPARATION)(.*?)(?=\n\n[A-Z]|$)',
                r'(?:General Procedures|GENERAL PROCEDURES)(.*?)(?=\n\n[A-Z]|$)',
            ]
            
            for pattern in exp_patterns:
                match = re.search(pattern, all_text, re.DOTALL | re.IGNORECASE)
                if match:
                    experimental_text = match.group(1).strip()
                    logger.info(f"  找到实验部分")
                    break
            
            # 如果没找到，查找包含合成关键词的段落
            if not experimental_text:
                logger.info("  未找到标准实验部分，搜索合成相关段落...")
                lines = all_text.split('\n')
                synth_lines = []
                for line in lines:
                    if re.search(r'(synthesized|prepared|compound|reaction|procedure|method|synthesis)', line, re.IGNORECASE):
                        synth_lines.append(line.strip())
                
                if synth_lines:
                    experimental_text = '\n'.join(synth_lines[:50])
                    logger.info(f"  找到 {len(synth_lines)} 行合成相关内容")
            
            # 提取化合物编号
            compounds = set()
            compound_patterns = [
                r'compound\s+(\d+|[A-Z]\d*)',
                r'Compounds?\s+(\d+|[A-Z]\d*)',
                r'(\d+|[A-Z]\d*)\s*\([^)]+\)',
                r'([A-Z]+-\d+)',
            ]
            
            for pattern in compound_patterns:
                matches = re.findall(pattern, all_text, re.IGNORECASE)
                for match in matches:
                    if isinstance(match, tuple):
                        match = match[0]
                    if len(match) > 1 or match.isdigit():
                        compounds.add(match)
            
            # 提取物理性质数据
            phys_props = []
            phys_patterns = [
                r'm\.p\.:.*?(?=\n\n|\n[A-Z])',
                r'Tdec.*?(?=\n\n|\n[A-Z])',
                r'Tpeak.*?(?=\n\n|\n[A-Z])',
                r'IR.*?(?=\n\n|\n[A-Z])',
                r'1H NMR.*?(?=\n\n|\n[A-Z])',
                r'13C NMR.*?(?=\n\n|\n[A-Z])',
                r'HRMS.*?(?=\n\n|\n[A-Z])',
            ]
            
            for pattern in phys_patterns:
                matches = re.findall(pattern, all_text, re.DOTALL | re.IGNORECASE)
                phys_props.extend(matches)
            
            # 提取化合物描述
            compound_desc = {}
            if compounds:
                for compound in list(compounds)[:5]:
                    pattern = fr'{re.escape(compound)}.*?(?=\n\n|\n[A-Z]|\.\s)'
                    match = re.search(pattern, all_text, re.DOTALL | re.IGNORECASE)
                    if match:
                        desc = match.group(0)[:200]
                        compound_desc[compound] = desc
            
            return {
                'experimental_text': experimental_text,
                'compounds': sorted(compounds, key=lambda x: (str(x).isnumeric(), str(x))),
                'physical_props': phys_props[:10],
                'compound_descriptions': compound_desc,
                'total_pages': total_pages,
                'pages_processed': pages_to_process,
                'source_pdf': str(pdf_path)
            }
            
    except Exception as e:
        logger.error(f"提取文本信息失败: {e}")
        return {
            'experimental_text': '',
            'compounds': [],
            'physical_props': [],
            'compound_descriptions': {},
            'total_pages': 0,
            'pages_processed': 0,
            'source_pdf': str(pdf_path),
            'error': str(e)
        }

def extract_reactions_with_openchemie(pdf_path: Path, max_pages: int = 5) -> List[Dict[str, Any]]:
    """使用OpenChemIE从PDF图表中提取化学反应"""
    logger.info(f"使用OpenChemIE提取反应: {pdf_path.name}")
    
    try:
        import torch
        # 添加OpenChemIE路径
        openchemie_path = Path(__file__).parent / "openchemie_test" / "OpenChemIE"
        if openchemie_path.exists():
            sys.path.insert(0, str(openchemie_path))
        
        from openchemie import OpenChemIE
        
        # 初始化OpenChemIE
        device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        logger.info(f"使用设备: {device}")
        
        model = OpenChemIE(device=device)
        
        # 提取反应
        figure_results = model.extract_reactions_from_figures_in_pdf(
            str(pdf_path),
            num_pages=max_pages,
            molscribe=True,
            ocr=False
        )
        
        # 处理结果
        reactions = []
        for fig_idx, fig_result in enumerate(figure_results):
            fig_reactions = fig_result.get('reactions', [])
            page = fig_result.get('page', fig_idx)
            
            for rxn_idx, rxn in enumerate(fig_reactions):
                # 提取SMILES
                reactant_smiles = None
                product_smiles = None
                
                if 'reactants' in rxn and rxn['reactants']:
                    for reactant in rxn['reactants']:
                        if 'smiles' in reactant and reactant['smiles']:
                            reactant_smiles = reactant['smiles']
                            break
                
                if 'products' in rxn and rxn['products']:
                    for product in rxn['products']:
                        if 'smiles' in product and product['smiles']:
                            product_smiles = product['smiles']
                            break
                
                if reactant_smiles and product_smiles:
                    reaction_data = {
                        'reaction_id': len(reactions) + 1,
                        'page': page + 1,
                        'reactant_smiles': reactant_smiles,
                        'product_smiles': product_smiles,
                        'raw_data': rxn
                    }
                    reactions.append(reaction_data)
                    logger.info(f"  提取到反应 {reaction_data['reaction_id']} (第{reaction_data['page']}页)")
        
        logger.info(f"总共提取到 {len(reactions)} 个反应")
        return reactions
        
    except Exception as e:
        logger.error(f"使用OpenChemIE提取反应失败: {e}")
        import traceback
        traceback.print_exc()
        return []

def generate_reaction_images(reactions: List[Dict[str, Any]], output_dir: Path) -> List[Dict[str, Any]]:
    """为提取的反应生成图像"""
    try:
        from rdkit import Chem
        from rdkit.Chem import Draw
        from PIL import Image
    except ImportError:
        logger.error("RDKit未安装，无法生成反应图像")
        return reactions
    
    logger.info("生成反应图像...")
    
    images_dir = output_dir / "reaction_images"
    images_dir.mkdir(exist_ok=True)
    
    for i, reaction in enumerate(reactions):
        reaction_num = reaction['reaction_id']
        
        try:
            # 解析SMILES
            reactant_mol = Chem.MolFromSmiles(reaction['reactant_smiles'])
            product_mol = Chem.MolFromSmiles(reaction['product_smiles'])
            
            if reactant_mol and product_mol:
                # 生成反应物图像
                reactant_img = Draw.MolToImage(reactant_mol, size=(300, 300))
                reactant_path = images_dir / f"reaction_{reaction_num}_reactant.png"
                reactant_img.save(reactant_path)
                
                # 生成产物图像
                product_img = Draw.MolToImage(product_mol, size=(300, 300))
                product_path = images_dir / f"reaction_{reaction_num}_product.png"
                product_img.save(product_path)
                
                # 更新反应数据
                reaction['reactant_image'] = str(reactant_path)
                reaction['product_image'] = str(product_path)
                reaction['images_generated'] = True
                
                logger.info(f"  反应 {reaction_num}: 生成2张图像")
            else:
                logger.warning(f"  反应 {reaction_num}: 无法解析SMILES")
                reaction['images_generated'] = False
                
        except Exception as e:
            logger.error(f"  反应 {reaction_num}: 生成图像失败 - {e}")
            reaction['images_generated'] = False
    
    return reactions

def create_word_document(pdf_info: Dict[str, Any], reactions: List[Dict[str, Any]], 
                         output_path: Path, language: str = 'both', 
                         template_path: Optional[Path] = None) -> bool:
    """创建Word文档"""
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx未安装，无法创建Word文档")
        return False
    
    logger.info(f"创建Word文档: {output_path.name}")
    
    # 加载模板或创建新文档
    if template_path and template_path.exists():
        doc = docx.Document(str(template_path))
    else:
        doc = docx.Document()
    
    # 添加标题
    if pdf_info['compounds']:
        title = f"化合物{', '.join(pdf_info['compounds'][:3])}的合成路线"
        if len(pdf_info['compounds']) > 3:
            title += "等"
    else:
        title = f"{Path(pdf_info['source_pdf']).stem} 合成路线"
    
    title_para = doc.add_paragraph(title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].bold = True
    title_para.runs[0].font.size = Pt(16)
    
    # 空行
    doc.add_paragraph()
    
    # 原文部分（英文）
    if language in ['en', 'both']:
        doc.add_paragraph("原文：")
        
        if pdf_info['experimental_text']:
            exp_text = pdf_info['experimental_text']
            # 如果太长，截断
            if len(exp_text) > 5000:
                exp_text = exp_text[:5000] + "\n\n...（内容过长，已截断）"
            
            # 格式化文本
            lines = exp_text.split('\n')
            for line in lines:
                if line.strip():
                    # 检查是否以化合物编号开头
                    compound_start = False
                    for compound in pdf_info['compounds']:
                        if line.strip().startswith(compound) or f"compound {compound}" in line.lower():
                            compound_start = True
                            break
                    
                    if compound_start:
                        # 化合物段落加粗
                        para = doc.add_paragraph()
                        run = para.add_run(line)
                        run.bold = True
                    else:
                        doc.add_paragraph(line)
        else:
            doc.add_paragraph("（实验部分未提取到）")
    
    # 物理性质数据
    if pdf_info['physical_props']:
        doc.add_paragraph()
        doc.add_paragraph("物理性质数据：")
        for prop in pdf_info['physical_props']:
            doc.add_paragraph(prop)
    
    # 中文部分
    if language in ['zh', 'both']:
        doc.add_paragraph()
        doc.add_paragraph("中文：")
        doc.add_paragraph("（待翻译）")
    
    # 反应图部分（如果有反应）
    if reactions:
        doc.add_paragraph()
        doc.add_paragraph("反应图：")
        explanation = doc.add_paragraph("以下是通过OpenChemIE自动提取的化学反应结构图：")
        
        for reaction in reactions:
            reaction_num = reaction['reaction_id']
            
            # 添加反应标题
            reaction_header = doc.add_paragraph(f"反应 {reaction_num} (第{reaction['page']}页):")
            reaction_header.runs[0].italic = True
            
            # 添加SMILES信息
            smiles_text = f"反应物: {reaction['reactant_smiles'][:50]}... → 产物: {reaction['product_smiles'][:50]}..."
            smiles_para = doc.add_paragraph(smiles_text)
            smiles_para.runs[0].font.size = Pt(9)
            
            # 插入图像（如果已生成）
            if reaction.get('images_generated', False):
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 插入反应物图像
                if 'reactant_image' in reaction and Path(reaction['reactant_image']).exists():
                    try:
                        run = img_para.add_run()
                        run.add_picture(reaction['reactant_image'], width=Inches(1.8))
                        run.add_text("   →   ")
                        logger.info(f"  插入反应物图像: {Path(reaction['reactant_image']).name}")
                    except Exception as e:
                        logger.error(f"  插入反应物图像失败: {e}")
                
                # 插入产物图像
                if 'product_image' in reaction and Path(reaction['product_image']).exists():
                    try:
                        run = img_para.add_run()
                        run.add_picture(reaction['product_image'], width=Inches(1.8))
                        logger.info(f"  插入产物图像: {Path(reaction['product_image']).name}")
                    except Exception as e:
                        logger.error(f"  插入产物图像失败: {e}")
            else:
                no_image = doc.add_paragraph("（反应图像未生成）")
                no_image.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            
            # 添加空行分隔
            doc.add_paragraph()
    
    # 谱图部分（占位符）
    doc.add_paragraph()
    doc.add_paragraph("谱图：")
    doc.add_paragraph("（待补充）")
    
    # 添加元数据
    doc.add_paragraph()
    doc.add_paragraph("-" * 50)
    meta_text = f"文档生成信息: 来源PDF - {Path(pdf_info['source_pdf']).name}, "
    meta_text += f"处理页数 - {pdf_info['pages_processed']}/{pdf_info['total_pages']}, "
    meta_text += f"提取反应数 - {len(reactions)}"
    meta_para = doc.add_paragraph(meta_text)
    meta_para.runs[0].font.size = Pt(8)
    meta_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存文档
    doc.save(str(output_path))
    logger.info(f"文档已保存: {output_path}")
    
    return True

def process_pdf(pdf_path: Path, output_dir: Path, args: argparse.Namespace, 
                config: Dict[str, Any]) -> Dict[str, Any]:
    """处理单个PDF文件"""
    logger.info(f"\n{'='*60}")
    logger.info(f"处理文件: {pdf_path.name}")
    logger.info('='*60)
    
    # 创建输出子目录
    pdf_output_dir = output_dir / pdf_path.stem
    pdf_output_dir.mkdir(exist_ok=True)
    
    result = {
        'pdf_file': str(pdf_path),
        'output_dir': str(pdf_output_dir),
        'success': False,
        'errors': []
    }
    
    try:
        # 1. 提取文本信息
        text_info = extract_text_info(pdf_path, args.max_pages)
        result['text_info'] = {
            'compounds_found': len(text_info['compounds']),
            'experimental_text_length': len(text_info['experimental_text']),
            'physical_props_count': len(text_info['physical_props'])
        }
        
        # 2. 使用OpenChemIE提取反应
        reactions = extract_reactions_with_openchemie(pdf_path, args.max_pages)
        result['reactions_extracted'] = len(reactions)
        
        # 3. 生成反应图像（如果不需要则跳过）
        if reactions and not args.no_images:
            reactions = generate_reaction_images(reactions, pdf_output_dir)
            result['images_generated'] = sum(1 for r in reactions if r.get('images_generated', False))
        
        # 4. 创建Word文档
        output_doc_path = pdf_output_dir / f"{pdf_path.stem}_合成路线.docx"
        
        if create_word_document(text_info, reactions, output_doc_path, args.language, 
                               Path(args.template) if args.template else None):
            result['output_doc'] = str(output_doc_path)
            result['success'] = True
            
            # 记录文件大小
            if output_doc_path.exists():
                result['doc_size_kb'] = output_doc_path.stat().st_size // 1024
        else:
            result['errors'].append("创建Word文档失败")
        
        # 5. 保存原始数据（可选）
        if config.get('save_raw_data', False):
            raw_data_path = pdf_output_dir / f"{pdf_path.stem}_raw_data.json"
            raw_data = {
                'text_info': text_info,
                'reactions': reactions,
                'processing_args': vars(args)
            }
            with open(raw_data_path, 'w', encoding='utf-8') as f:
                json.dump(raw_data, f, ensure_ascii=False, indent=2)
            result['raw_data_file'] = str(raw_data_path)
        
    except Exception as e:
        logger.error(f"处理PDF时发生错误: {e}")
        result['errors'].append(str(e))
        import traceback
        traceback.print_exc()
    
    return result

def main():
    """主函数"""
    parser = setup_argparse()
    args = parser.parse_args()
    
    # 设置日志级别
    if args.debug:
        logging.getLogger().setLevel(logging.DEBUG)
        logger.debug("调试模式已启用")
    
    # 加载配置文件
    config = {}
    if args.config:
        config = load_config(args.config)
    
    # 创建输出目录
    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 查找PDF文件
    pdf_files = find_pdf_files(args.input, args.include_si)
    
    if not pdf_files:
        logger.error("未找到PDF文件")
        return 1
    
    logger.info(f"找到 {len(pdf_files)} 个PDF文件")
    for pdf_file in pdf_files:
        logger.info(f"  - {pdf_file}")
    
    # 处理每个PDF文件
    results = []
    for pdf_file in pdf_files:
        result = process_pdf(pdf_file, output_dir, args, config)
        results.append(result)
    
    # 生成报告
    logger.info(f"\n{'='*60}")
    logger.info("处理完成!")
    logger.info('='*60)
    
    success_count = sum(1 for r in results if r['success'])
    total_reactions = sum(r.get('reactions_extracted', 0) for r in results)
    
    logger.info(f"成功处理: {success_count}/{len(results)} 个文件")
    logger.info(f"总共提取到: {total_reactions} 个反应")
    
    # 保存处理报告
    report_path = output_dir / "processing_report.json"
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump({
            'timestamp': str(pd.Timestamp.now()) if HAS_PANDAS else str(datetime.now()),
            'args': vars(args),
            'results': results
        }, f, ensure_ascii=False, indent=2)
    
    logger.info(f"详细报告已保存: {report_path}")
    
    # 显示简要结果
    logger.info("\n处理结果摘要:")
    for result in results:
        pdf_name = Path(result['pdf_file']).name
        status = "✅ 成功" if result['success'] else "❌ 失败"
        logger.info(f"  {pdf_name}: {status}")
        if result['success']:
            logger.info(f"      化合物: {result['text_info']['compounds_found']} 个")
            logger.info(f"      反应: {result.get('reactions_extracted', 0)} 个")
            logger.info(f"      图像: {result.get('images_generated', 0)} 个")
            if 'output_doc' in result:
                logger.info(f"      文档: {Path(result['output_doc']).name} ({result.get('doc_size_kb', 0)} KB)")
    
    return 0 if success_count > 0 else 1

if __name__ == "__main__":
    sys.exit(main())